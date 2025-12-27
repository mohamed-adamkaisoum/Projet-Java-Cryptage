# -*- coding: utf-8 -*-
"""
Script de conversion Markdown (.md) vers Word (.docx) pour documentation Java
Utilisation :
    1. Placez "Documentation-Cryptage-Java.md" à côté de ce script.
    2. Installez python-docx :
        pip install python-docx
    3. Lancez :
        python md2docx.py
    -> Le fichier 'Documentation-Cryptage-Java.docx' sera généré.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Personnalisation : couleurs, polices, tailles
title_color = RGBColor(0x1a, 0x23, 0x7e)    # Bleu foncé
code_back_color = RGBColor(0xec, 0xef, 0xf1) # Gris très clair
code_font_color = RGBColor(0x26, 0x32, 0x38) # Gris code


def style_heading(para, level):
    run = para.runs[0]
    run.font.size = Pt(max(22 - 2*level, 11))
    run.font.bold = True
    run.font.color.rgb = title_color
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_codeblock(doc, code_lines):
    code = '\n'.join(code_lines)
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10.5)
    run.font.color.rgb = code_font_color
    shading_elm = para._element.get_or_add_pPr().add_new_shd()
    shading_elm.val = 'clear'
    shading_elm.color = 'auto'
    shading_elm.fill = 'eceff1'
    para.style = 'No Spacing'


def convert_markdown_to_docx(md_file, docx_file):
    doc = Document()
    with open(md_file, encoding='utf-8') as f:
        in_code = False
        code_lines = []
        for line in f:
            stripped = line.rstrip('\n')
            if stripped.startswith('#') and not in_code:
                level = len(stripped) - len(stripped.lstrip('#'))
                title = stripped.lstrip('#').strip()
                para = doc.add_paragraph(title)
                style_heading(para, level)
            elif stripped.startswith('```') and not in_code:
                in_code = True
                code_lines = []
            elif stripped.startswith('```') and in_code:
                add_codeblock(doc, code_lines)
                in_code = False
            elif in_code:
                code_lines.append(stripped)
            elif stripped == '' or stripped.isspace():
                doc.add_paragraph('')
            else:
                doc.add_paragraph(stripped)
    doc.save(docx_file)

if __name__ == "__main__":
    convert_markdown_to_docx("Documentation-Cryptage-Java.md", "Documentation-Cryptage-Java.docx")
    print("✅ Fichier Word généré : Documentation-Cryptage-Java.docx")

